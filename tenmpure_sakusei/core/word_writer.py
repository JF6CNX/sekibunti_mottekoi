from docx import Document
from docx.oxml.ns import qn


# ===== フォント設定 =====
def set_font(run):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


# ===== CDCl₃専用（ここだけ下付き） =====
def add_cdcl3(paragraph, bold=False):
    run1 = paragraph.add_run("CDCl")
    run1.bold = bold
    set_font(run1)

    run2 = paragraph.add_run("3")
    run2.font.subscript = True
    run2.bold = bold
    set_font(run2)


# ===== NMR行 =====
def add_nmr_line(doc, text, nucleus_label):

    p = doc.add_paragraph()

    # ===== 核種 =====
    number = "".join([c for c in nucleus_label if c.isdigit()])
    symbol = "".join([c for c in nucleus_label if not c.isdigit()])

    run_num = p.add_run(number)
    run_num.font.superscript = True
    run_num.bold = True
    set_font(run_num)

    run_symbol = p.add_run(symbol)
    run_symbol.bold = True
    set_font(run_symbol)

    # ===== 本文 =====
    rest = text.replace(f"{nucleus_label} ", "", 1)

    if "δ" in rest:
        before, after = rest.split("δ", 1)

        # ===== NMRヘッダ（太字） =====
        if "CDCl3" in before:
            pre, post = before.split("CDCl3")

            run_pre = p.add_run(" " + pre)
            run_pre.bold = True
            set_font(run_pre)

            add_cdcl3(p, bold=True)

            run_post = p.add_run(post)
            run_post.bold = True
            set_font(run_post)

        else:
            run_before = p.add_run(" " + before)
            run_before.bold = True
            set_font(run_before)

        # ===== δ（イタリック） =====
        run_delta = p.add_run("δ")
        run_delta.italic = True
        set_font(run_delta)

        # ===== データ（←絶対に加工しない） =====
        run_after = p.add_run(after)
        set_font(run_after)

    else:
        run_rest = p.add_run(" " + rest)
        set_font(run_rest)


# ===== Word出力 =====
def write_word(template, output_path):

    doc = Document()

    # タイトル
    p = doc.add_paragraph()
    run = p.add_run(f"[{template['id']}]")
    set_font(run)

    # 構造
    if template["structure"]:
        p = doc.add_paragraph()
        run = p.add_run(template["structure"])
        set_font(run)

    # 各NMR
    if template["H1"]:
        add_nmr_line(doc, template["H1"], "1H")

    if template["C13"]:
        add_nmr_line(doc, template["C13"], "13C")

    if template["F19"]:
        add_nmr_line(doc, template["F19"], "19F")

    doc.save(output_path)